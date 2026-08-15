"""
Asset document parsing, storage and versioning service.
"""

from __future__ import annotations

import io
import html
import mimetypes
import os
import re
import zipfile
import xml.etree.ElementTree as ET
from collections import defaultdict
from typing import Any, Dict, List, Optional, Tuple

from sqlalchemy import func
from sqlalchemy.orm import Session

from app.domains.asset.models.asset import AssetType, SddAsset, SddAssetVersion
from app.domains.task.models.task import SddTask

try:
    from docx import Document as DocxDocument  # type: ignore
except Exception:
    DocxDocument = None


SUPPORTED_INLINE_REVIEW_EXTENSIONS = {".md", ".markdown", ".txt", ".docx"}
DOCX_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
DOCX_NS_ATTR = f"{{{DOCX_NS}}}"
DOCX_NS_MAP = {"w": DOCX_NS}


def _task_assets_root(task: SddTask) -> str:
    raw_project_path = str(task.project_path or "").strip()
    if not raw_project_path:
        raise ValueError("Task project path is missing")
    root = os.path.abspath(os.path.join(raw_project_path, ".sdd", "assets"))
    os.makedirs(root, exist_ok=True)
    return root


def _task_for_asset(db: Session, asset: SddAsset) -> SddTask:
    task = db.query(SddTask).filter(SddTask.id == asset.task_id).first()
    if not task:
        raise ValueError(f"Task not found for asset {asset.id}")
    return task


def _normalize_filename(file_name: Optional[str], fallback_ext: str = ".md") -> str:
    base = os.path.basename(file_name or "").strip()
    if not base:
        base = f"document{fallback_ext}"
    sanitized = re.sub(r"[^A-Za-z0-9._-]+", "_", base).strip("._")
    return sanitized or f"document{fallback_ext}"


def _guess_ext_and_mime(file_name: Optional[str]) -> Tuple[str, str]:
    name = file_name or ""
    ext = os.path.splitext(name)[1].lower()
    mime = mimetypes.guess_type(name)[0] or "application/octet-stream"
    if ext == ".md" or ext == ".markdown":
        mime = "text/markdown"
    elif ext == ".txt":
        mime = "text/plain"
    elif ext == ".docx":
        mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    elif ext == ".pdf":
        mime = "application/pdf"
    return ext, mime


def _decode_text_bytes(raw: bytes) -> str:
    if not raw:
        return ""
    for encoding in ("utf-8", "utf-8-sig", "gb18030", "latin-1"):
        try:
            return raw.decode(encoding)
        except Exception:
            continue
    return raw.decode("utf-8", errors="ignore")


def _xml_attr(name: str) -> str:
    return f"{DOCX_NS_ATTR}{name}"


def _local_name(tag: str) -> str:
    if not tag:
        return ""
    if "}" in tag:
        return tag.rsplit("}", 1)[-1]
    return tag


def _safe_int(value: Any, default: int = 0) -> int:
    try:
        return int(value)
    except Exception:
        return default


def _merge_adjacent_runs(runs: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    if not runs:
        return []
    merged: List[Dict[str, Any]] = []
    for run in runs:
        text = str(run.get("text") or "")
        if not text:
            continue
        style = {k: v for k, v in run.items() if k != "text"}
        if merged:
            prev = merged[-1]
            prev_style = {k: v for k, v in prev.items() if k != "text"}
            if prev_style == style:
                prev["text"] = str(prev.get("text") or "") + text
                continue
        merged.append({"text": text, **style})
    return merged


def _render_runs_to_text(runs: List[Dict[str, Any]]) -> str:
    return "".join(str(item.get("text") or "") for item in runs)


def _word_bool(node: Optional[ET.Element]) -> bool:
    if node is None:
        return False
    raw = str(node.attrib.get(_xml_attr("val"), "") or "").strip().lower()
    if raw in {"0", "false", "off", "none"}:
        return False
    return True


def _word_highlight_to_hex(name: str) -> Optional[str]:
    palette = {
        "yellow": "#fef08a",
        "green": "#86efac",
        "cyan": "#67e8f9",
        "magenta": "#f9a8d4",
        "blue": "#93c5fd",
        "red": "#fca5a5",
        "darkblue": "#1d4ed8",
        "darkred": "#b91c1c",
        "darkyellow": "#ca8a04",
        "darkgreen": "#15803d",
        "darkcyan": "#0e7490",
        "darkmagenta": "#a21caf",
        "lightgray": "#e5e7eb",
        "darkgray": "#6b7280",
        "black": "#111827",
    }
    return palette.get(name.strip().lower())


def _hex_color(raw: str) -> Optional[str]:
    normalized = raw.strip().replace("#", "").upper()
    if not normalized or normalized in {"AUTO", "NONE"}:
        return None
    if re.fullmatch(r"[0-9A-F]{6}", normalized):
        return f"#{normalized}"
    return None


def _decode_docx_symbol(char_hex: str) -> str:
    raw = (char_hex or "").strip().upper()
    if not raw:
        return ""
    common = {
        "F0B7": "•",
        "00B7": "•",
        "F0A7": "§",
        "F0D8": "°",
        "F0FC": "✓",
        "F0A8": "→",
        "F0E8": "➢",
    }
    if raw in common:
        return common[raw]
    try:
        value = int(raw, 16)
    except Exception:
        return ""
    if value < 0 or value > 0x10FFFF:
        return ""
    try:
        return chr(value)
    except Exception:
        return ""


def _normalize_list_marker(marker: str) -> str:
    text = (marker or "").strip()
    if not text:
        return "•"
    for ch in text:
        code = ord(ch)
        if 0xF000 <= code <= 0xF8FF:
            return "•"
    return text


def _roman_numeral(value: int, uppercase: bool = True) -> str:
    if value <= 0:
        return str(value)
    numerals = [
        (1000, "M"),
        (900, "CM"),
        (500, "D"),
        (400, "CD"),
        (100, "C"),
        (90, "XC"),
        (50, "L"),
        (40, "XL"),
        (10, "X"),
        (9, "IX"),
        (5, "V"),
        (4, "IV"),
        (1, "I"),
    ]
    out = []
    remaining = value
    for unit, symbol in numerals:
        while remaining >= unit:
            out.append(symbol)
            remaining -= unit
    rendered = "".join(out) or str(value)
    return rendered if uppercase else rendered.lower()


def _alpha_numeral(value: int, uppercase: bool = False) -> str:
    if value <= 0:
        return str(value)
    chars = []
    current = value
    while current > 0:
        current -= 1
        chars.append(chr((current % 26) + (65 if uppercase else 97)))
        current //= 26
    rendered = "".join(reversed(chars))
    return rendered.upper() if uppercase else rendered.lower()


def _markdown_to_blocks(markdown: str) -> List[Dict[str, Any]]:
    lines = markdown.replace("\r\n", "\n").replace("\r", "\n").split("\n")
    blocks: List[Dict[str, Any]] = []
    current_para: List[str] = []
    block_index = 0

    def flush_paragraph() -> None:
        nonlocal block_index
        if not current_para:
            return
        text = " ".join(part.strip() for part in current_para if part.strip()).strip()
        current_para.clear()
        if not text:
            return
        block_index += 1
        blocks.append(
            {
                "id": f"blk-{block_index}",
                "type": "paragraph",
                "text": text,
                "order": block_index,
                "meta": {},
            }
        )

    for line in lines:
        stripped = line.strip()
        if not stripped:
            flush_paragraph()
            continue

        heading_match = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading_match:
            flush_paragraph()
            block_index += 1
            level = len(heading_match.group(1))
            blocks.append(
                {
                    "id": f"blk-{block_index}",
                    "type": "heading",
                    "text": heading_match.group(2).strip(),
                    "order": block_index,
                    "meta": {"level": level},
                }
            )
            continue

        list_match = re.match(r"^([-*+]|\d+\.)\s+(.+)$", stripped)
        if list_match:
            flush_paragraph()
            block_index += 1
            marker = list_match.group(1)
            blocks.append(
                {
                    "id": f"blk-{block_index}",
                    "type": "list_item",
                    "text": list_match.group(2).strip(),
                    "order": block_index,
                    "meta": {"marker": marker},
                }
            )
            continue

        current_para.append(stripped)

    flush_paragraph()
    return blocks


def _plain_text_to_markdown(text: str) -> str:
    lines = text.replace("\r\n", "\n").replace("\r", "\n").split("\n")
    out: List[str] = []
    buffer: List[str] = []

    def flush() -> None:
        if not buffer:
            return
        out.append(" ".join(part.strip() for part in buffer if part.strip()).strip())
        buffer.clear()

    for line in lines:
        stripped = line.strip()
        if not stripped:
            flush()
            if out and out[-1] != "":
                out.append("")
            continue
        buffer.append(stripped)

    flush()
    while out and out[-1] == "":
        out.pop()

    return "\n".join(out).strip()


def _extract_docx_comment_text(node: ET.Element) -> str:
    lines: List[str] = []
    for paragraph in node.findall("./w:p", DOCX_NS_MAP):
        parts: List[str] = []
        for item in paragraph.iter():
            lname = _local_name(item.tag)
            if lname == "t" and item.text:
                parts.append(item.text)
            elif lname in {"tab", "ptab"}:
                parts.append("\t")
            elif lname in {"br", "cr"}:
                parts.append("\n")
            elif lname == "sym":
                parts.append(_decode_docx_symbol(item.attrib.get(_xml_attr("char"), "")))
        line = "".join(parts).strip()
        if line:
            lines.append(line)
    return "\n".join(lines).strip()


def _parse_docx_comments(comments_root: Optional[ET.Element]) -> Dict[str, Dict[str, Any]]:
    comments: Dict[str, Dict[str, Any]] = {}
    if comments_root is None:
        return comments

    for comment in comments_root.findall(".//w:comment", DOCX_NS_MAP):
        comment_id = str(comment.attrib.get(_xml_attr("id"), "") or "").strip()
        if not comment_id:
            continue
        comments[comment_id] = {
            "comment_id": comment_id,
            "author": str(comment.attrib.get(_xml_attr("author"), "") or "").strip(),
            "initials": str(comment.attrib.get(_xml_attr("initials"), "") or "").strip(),
            "date": str(comment.attrib.get(_xml_attr("date"), "") or "").strip(),
            "content": _extract_docx_comment_text(comment),
        }
    return comments


def _parse_docx_numbering(numbering_root: Optional[ET.Element]) -> Dict[str, Any]:
    result: Dict[str, Any] = {"num_to_abs": {}, "levels": {}}
    if numbering_root is None:
        return result

    levels_map: Dict[str, Dict[int, Dict[str, Any]]] = {}
    for abstract in numbering_root.findall(".//w:abstractNum", DOCX_NS_MAP):
        abstract_id = str(abstract.attrib.get(_xml_attr("abstractNumId"), "") or "").strip()
        if not abstract_id:
            continue
        level_info: Dict[int, Dict[str, Any]] = {}
        for level_node in abstract.findall("./w:lvl", DOCX_NS_MAP):
            level = _safe_int(level_node.attrib.get(_xml_attr("ilvl")), 0)
            num_fmt_node = level_node.find("./w:numFmt", DOCX_NS_MAP)
            lvl_text_node = level_node.find("./w:lvlText", DOCX_NS_MAP)
            start_node = level_node.find("./w:start", DOCX_NS_MAP)
            level_info[level] = {
                "num_fmt": str(num_fmt_node.attrib.get(_xml_attr("val"), "") if num_fmt_node is not None else "").strip().lower(),
                "lvl_text": str(lvl_text_node.attrib.get(_xml_attr("val"), "") if lvl_text_node is not None else "").strip(),
                "start": max(1, _safe_int(start_node.attrib.get(_xml_attr("val")) if start_node is not None else 1, 1)),
            }
        levels_map[abstract_id] = level_info

    num_to_abs: Dict[str, str] = {}
    for num_node in numbering_root.findall(".//w:num", DOCX_NS_MAP):
        num_id = str(num_node.attrib.get(_xml_attr("numId"), "") or "").strip()
        abs_node = num_node.find("./w:abstractNumId", DOCX_NS_MAP)
        abs_id = str(abs_node.attrib.get(_xml_attr("val"), "") if abs_node is not None else "").strip()
        if num_id and abs_id:
            num_to_abs[num_id] = abs_id

    result["num_to_abs"] = num_to_abs
    result["levels"] = levels_map
    return result


def _format_counter(value: int, num_fmt: str) -> str:
    fmt = (num_fmt or "decimal").lower()
    if fmt in {"decimal", "decimalzero"}:
        return str(value)
    if fmt in {"lowerletter", "loweralpha"}:
        return _alpha_numeral(value, uppercase=False)
    if fmt in {"upperletter", "upperalpha"}:
        return _alpha_numeral(value, uppercase=True)
    if fmt == "lowerroman":
        return _roman_numeral(value, uppercase=False)
    if fmt == "upperroman":
        return _roman_numeral(value, uppercase=True)
    return str(value)


def _resolve_list_marker(
    *,
    num_id: Optional[str],
    level: int,
    numbering: Dict[str, Any],
    numbering_state: Dict[str, List[int]],
) -> str:
    if not num_id:
        return "•"

    levels = numbering.get("levels") or {}
    num_to_abs = numbering.get("num_to_abs") or {}
    abstract_id = str(num_to_abs.get(num_id) or "").strip()
    level_info = (levels.get(abstract_id) or {}).get(level) or {}

    start_value = max(1, _safe_int(level_info.get("start"), 1))
    counters = numbering_state[num_id]
    while len(counters) <= level:
        counters.append(0)
    if counters[level] <= 0:
        counters[level] = start_value
    else:
        counters[level] += 1
    for idx in range(level + 1, len(counters)):
        counters[idx] = 0

    num_fmt = str(level_info.get("num_fmt") or "").strip().lower()
    lvl_text = str(level_info.get("lvl_text") or "").strip()
    if num_fmt == "bullet":
        bullet = re.sub(r"%\d+", "", lvl_text).strip() or "•"
        bullet = _normalize_list_marker(bullet)
        return bullet

    template = lvl_text or f"%{level + 1}."

    def repl(match: re.Match[str]) -> str:
        requested_idx = max(0, _safe_int(match.group(1), level + 1) - 1)
        if requested_idx >= len(counters) or counters[requested_idx] <= 0:
            value = 1
        else:
            value = counters[requested_idx]
        if requested_idx == level:
            return _format_counter(value, num_fmt)
        return str(value)

    marker = re.sub(r"%(\d+)", repl, template).strip()
    if marker:
        return _normalize_list_marker(marker)
    return f"{_format_counter(counters[level], num_fmt)}."


def _extract_run_style(run: ET.Element) -> Dict[str, Any]:
    style: Dict[str, Any] = {}
    r_pr = run.find("./w:rPr", DOCX_NS_MAP)
    if r_pr is None:
        return style

    if _word_bool(r_pr.find("./w:b", DOCX_NS_MAP)):
        style["bold"] = True
    if _word_bool(r_pr.find("./w:i", DOCX_NS_MAP)):
        style["italic"] = True

    underline = r_pr.find("./w:u", DOCX_NS_MAP)
    if underline is not None:
        raw = str(underline.attrib.get(_xml_attr("val"), "") or "").strip().lower()
        if raw not in {"none", "0", "false", "off"}:
            style["underline"] = True

    if _word_bool(r_pr.find("./w:strike", DOCX_NS_MAP)) or _word_bool(r_pr.find("./w:dstrike", DOCX_NS_MAP)):
        style["strike"] = True

    vert_align = r_pr.find("./w:vertAlign", DOCX_NS_MAP)
    if vert_align is not None:
        align = str(vert_align.attrib.get(_xml_attr("val"), "") or "").strip().lower()
        if align in {"superscript", "subscript"}:
            style[align] = True

    color_node = r_pr.find("./w:color", DOCX_NS_MAP)
    if color_node is not None:
        color = _hex_color(str(color_node.attrib.get(_xml_attr("val"), "") or ""))
        if color:
            style["color"] = color

    highlight_node = r_pr.find("./w:highlight", DOCX_NS_MAP)
    if highlight_node is not None:
        highlight = _word_highlight_to_hex(str(highlight_node.attrib.get(_xml_attr("val"), "") or ""))
        if highlight:
            style["highlight"] = highlight

    size_node = r_pr.find("./w:sz", DOCX_NS_MAP)
    if size_node is not None:
        half_points = _safe_int(size_node.attrib.get(_xml_attr("val")), 0)
        if half_points > 0:
            style["font_size"] = round(half_points / 2, 2)

    font_node = r_pr.find("./w:rFonts", DOCX_NS_MAP)
    if font_node is not None:
        font_name = (
            str(font_node.attrib.get(_xml_attr("ascii"), "") or "").strip()
            or str(font_node.attrib.get(_xml_attr("hAnsi"), "") or "").strip()
            or str(font_node.attrib.get(_xml_attr("eastAsia"), "") or "").strip()
        )
        if font_name:
            style["font_name"] = font_name

    return style


def _extract_run_text(run: ET.Element) -> str:
    parts: List[str] = []
    for node in list(run):
        lname = _local_name(node.tag)
        if lname == "t" and node.text:
            parts.append(node.text)
        elif lname in {"tab", "ptab"}:
            parts.append("\t")
        elif lname in {"br", "cr"}:
            parts.append("\n")
        elif lname in {"noBreakHyphen", "softHyphen"}:
            parts.append("-")
        elif lname == "sym":
            parts.append(_decode_docx_symbol(node.attrib.get(_xml_attr("char"), "")))
        elif lname == "instrText" and node.text:
            parts.append(node.text)
    return "".join(parts)


def _extract_paragraph_meta(paragraph: ET.Element) -> Dict[str, Any]:
    meta: Dict[str, Any] = {}
    p_pr = paragraph.find("./w:pPr", DOCX_NS_MAP)
    if p_pr is None:
        return meta

    style_node = p_pr.find("./w:pStyle", DOCX_NS_MAP)
    if style_node is not None:
        style_val = str(style_node.attrib.get(_xml_attr("val"), "") or "").strip()
        if style_val:
            meta["style"] = style_val

    num_pr = p_pr.find("./w:numPr", DOCX_NS_MAP)
    if num_pr is not None:
        num_id_node = num_pr.find("./w:numId", DOCX_NS_MAP)
        ilvl_node = num_pr.find("./w:ilvl", DOCX_NS_MAP)
        num_id = str(num_id_node.attrib.get(_xml_attr("val"), "") if num_id_node is not None else "").strip()
        ilvl = max(0, _safe_int(ilvl_node.attrib.get(_xml_attr("val")) if ilvl_node is not None else 0, 0))
        if num_id:
            meta["num_id"] = num_id
        meta["level"] = ilvl

    return meta


def _parse_paragraph_content(
    paragraph: ET.Element,
    comments_by_id: Dict[str, Dict[str, Any]],
) -> Dict[str, Any]:
    runs: List[Dict[str, Any]] = []
    comment_ranges: Dict[str, Dict[str, Optional[int]]] = {}
    active_comment_ids: List[str] = []
    cursor = 0

    def ensure_comment(comment_id: str) -> Dict[str, Optional[int]]:
        if comment_id not in comment_ranges:
            comment_ranges[comment_id] = {"char_start": None, "char_end": None}
        return comment_ranges[comment_id]

    def start_comment(comment_id: str) -> None:
        nonlocal cursor
        if not comment_id:
            return
        info = ensure_comment(comment_id)
        if info["char_start"] is None:
            info["char_start"] = cursor
        if comment_id not in active_comment_ids:
            active_comment_ids.append(comment_id)

    def end_comment(comment_id: str) -> None:
        nonlocal cursor
        if not comment_id:
            return
        info = ensure_comment(comment_id)
        if info["char_start"] is None:
            info["char_start"] = cursor
        info["char_end"] = cursor
        if comment_id in active_comment_ids:
            active_comment_ids.remove(comment_id)

    def append_text(text: str, style: Dict[str, Any]) -> None:
        nonlocal cursor
        if not text:
            return
        run = {"text": text}
        run.update(style)
        runs.append(run)
        next_cursor = cursor + len(text)
        for comment_id in list(active_comment_ids):
            info = ensure_comment(comment_id)
            if info["char_start"] is None:
                info["char_start"] = cursor
            info["char_end"] = next_cursor
        cursor = next_cursor

    def walk(node: ET.Element) -> None:
        for child in list(node):
            lname = _local_name(child.tag)
            if lname == "commentRangeStart":
                start_comment(str(child.attrib.get(_xml_attr("id"), "") or "").strip())
                continue
            if lname == "commentRangeEnd":
                end_comment(str(child.attrib.get(_xml_attr("id"), "") or "").strip())
                continue
            if lname == "r":
                for ref in child.findall(".//w:commentReference", DOCX_NS_MAP):
                    ref_id = str(ref.attrib.get(_xml_attr("id"), "") or "").strip()
                    if ref_id:
                        info = ensure_comment(ref_id)
                        if info["char_start"] is None:
                            info["char_start"] = cursor
                        info["char_end"] = cursor
                append_text(_extract_run_text(child), _extract_run_style(child))
                continue
            if lname in {"hyperlink", "smartTag", "sdt", "ins", "del", "fldSimple"}:
                walk(child)
                continue
            if lname == "instrText" and child.text:
                append_text(child.text, {})

    walk(paragraph)
    for comment_id in list(active_comment_ids):
        end_comment(comment_id)

    merged_runs = _merge_adjacent_runs(runs)
    text = _render_runs_to_text(merged_runs)
    anchors: List[Dict[str, Any]] = []
    for comment_id, span in comment_ranges.items():
        start = span.get("char_start")
        end = span.get("char_end")
        if start is None and end is None:
            continue
        selected_text = ""
        if isinstance(start, int) and isinstance(end, int) and 0 <= start <= end <= len(text):
            selected_text = text[start:end]
        elif isinstance(start, int) and 0 <= start < len(text):
            selected_text = text[start:]
        if not selected_text:
            selected_text = str(comments_by_id.get(comment_id, {}).get("content") or "").split("\n", 1)[0].strip()
        anchors.append(
            {
                "comment_id": comment_id,
                "char_start": start,
                "char_end": end,
                "selected_text": selected_text or None,
            }
        )
    return {"text": text, "runs": merged_runs, "comment_anchors": anchors}


def _classify_paragraph_type(style_name: str, has_numbering: bool) -> Tuple[str, Dict[str, Any]]:
    style = (style_name or "").strip()
    style_l = style.lower()
    meta: Dict[str, Any] = {}
    if style:
        meta["style"] = style

    if style_l.startswith("heading"):
        level_match = re.search(r"(\d+)", style_l)
        level = min(max(_safe_int(level_match.group(1) if level_match else 1, 1), 1), 6)
        meta["level"] = level
        return "heading", meta

    if style_l in {"title"}:
        meta["level"] = 1
        return "heading", meta
    if style_l in {"subtitle"}:
        meta["level"] = 2
        return "heading", meta

    if "toc" in style_l:
        level_match = re.search(r"(\d+)", style_l)
        level = min(max(_safe_int(level_match.group(1) if level_match else 1, 1), 1), 6)
        meta["level"] = level
        return "toc_entry", meta

    if has_numbering or "list" in style_l:
        return "list_item", meta

    return "paragraph", meta


def _table_to_markdown(table_rows: List[Dict[str, Any]]) -> str:
    if not table_rows:
        return ""

    rows: List[List[str]] = []
    max_cols = 0
    for row in table_rows:
        cols = [str(cell.get("text") or "").replace("\n", "<br>").strip() for cell in (row.get("cells") or [])]
        rows.append(cols)
        max_cols = max(max_cols, len(cols))

    if max_cols <= 1:
        return "\n".join((item[0] if item else "") for item in rows).strip()

    padded_rows = [cols + [""] * (max_cols - len(cols)) for cols in rows]
    header = padded_rows[0]
    divider = ["---"] * max_cols
    lines = [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join(divider) + " |",
    ]
    for row in padded_rows[1:]:
        lines.append("| " + " | ".join(row) + " |")
    return "\n".join(lines).strip()


def _block_to_markdown_line(block: Dict[str, Any]) -> str:
    block_type = str(block.get("type") or "paragraph")
    text = str(block.get("text") or "")
    stripped = text.strip()
    meta = block.get("meta") or {}

    if block_type == "heading":
        level = min(max(_safe_int(meta.get("level"), 1), 1), 6)
        return f"{'#' * level} {stripped}".strip()

    if block_type == "list_item":
        marker = str(meta.get("marker") or "•").strip() or "•"
        level = max(0, _safe_int(meta.get("level"), 0))
        indent = "  " * level
        return f"{indent}{marker} {stripped}".rstrip()

    if block_type == "table":
        return _table_to_markdown(list(block.get("table", {}).get("rows") or []))

    return stripped


def _empty_docx_payload() -> Dict[str, Any]:
    return {
        "normalized_markdown": "",
        "blocks_json": [],
        "render_json": {
            "format": "rich_doc",
            "block_count": 0,
            "docx_comments": [],
            "features": {
                "runs": False,
                "tables": False,
                "toc": False,
                "comments": False,
            },
        },
    }


def _parse_docx_payload_via_xml(raw: bytes) -> Dict[str, Any]:
    try:
        with zipfile.ZipFile(io.BytesIO(raw)) as zf:
            document_xml = zf.read("word/document.xml")
            comments_xml = zf.read("word/comments.xml") if "word/comments.xml" in zf.namelist() else b""
            numbering_xml = zf.read("word/numbering.xml") if "word/numbering.xml" in zf.namelist() else b""
    except Exception:
        return _empty_docx_payload()

    try:
        document_root = ET.fromstring(document_xml)
    except Exception:
        return _empty_docx_payload()

    comments_root = None
    if comments_xml:
        try:
            comments_root = ET.fromstring(comments_xml)
        except Exception:
            comments_root = None

    numbering_root = None
    if numbering_xml:
        try:
            numbering_root = ET.fromstring(numbering_xml)
        except Exception:
            numbering_root = None

    comments_by_id = _parse_docx_comments(comments_root)
    numbering = _parse_docx_numbering(numbering_root)
    numbering_state: Dict[str, List[int]] = defaultdict(list)

    body = document_root.find("./w:body", DOCX_NS_MAP)
    if body is None:
        return _empty_docx_payload()

    blocks: List[Dict[str, Any]] = []
    markdown_lines: List[str] = []
    anchor_items: List[Dict[str, Any]] = []
    order = 0

    def append_block(block: Dict[str, Any], anchors: List[Dict[str, Any]]) -> None:
        nonlocal order
        order += 1
        block_id = f"blk-{order}"
        block["id"] = block_id
        block["order"] = order
        blocks.append(block)
        markdown_line = _block_to_markdown_line(block)
        if markdown_line:
            markdown_lines.append(markdown_line)
        for item in anchors:
            item["block_id"] = block_id
            anchor_items.append(item)

    for child in list(body):
        tag = _local_name(child.tag)
        if tag == "p":
            paragraph_data = _parse_paragraph_content(child, comments_by_id)
            text = str(paragraph_data.get("text") or "")
            if not text.strip():
                continue

            para_meta = _extract_paragraph_meta(child)
            style_name = str(para_meta.get("style") or "")
            has_numbering = bool(para_meta.get("num_id"))
            block_type, block_meta = _classify_paragraph_type(style_name, has_numbering)
            block_meta.update({k: v for k, v in para_meta.items() if k in {"style", "level"}})

            if block_type == "list_item":
                level = max(0, _safe_int(para_meta.get("level"), 0))
                num_id = str(para_meta.get("num_id") or "").strip() or None
                marker = _resolve_list_marker(
                    num_id=num_id,
                    level=level,
                    numbering=numbering,
                    numbering_state=numbering_state,
                )
                block_meta["level"] = level
                block_meta["marker"] = marker
                if num_id:
                    block_meta["num_id"] = num_id

            append_block(
                {
                    "type": block_type,
                    "text": text,
                    "runs": paragraph_data.get("runs") or [],
                    "meta": block_meta,
                },
                list(paragraph_data.get("comment_anchors") or []),
            )
            continue

        if tag == "tbl":
            rows_payload: List[Dict[str, Any]] = []
            row_line_texts: List[str] = []
            table_anchors: List[Dict[str, Any]] = []
            table_offset = 0

            for row in child.findall("./w:tr", DOCX_NS_MAP):
                row_cells: List[Dict[str, Any]] = []
                row_text_parts: List[str] = []
                row_anchor_items: List[Dict[str, Any]] = []
                row_cursor = 0

                for cell in row.findall("./w:tc", DOCX_NS_MAP):
                    cell_runs: List[Dict[str, Any]] = []
                    cell_text_parts: List[str] = []
                    cell_anchors: List[Dict[str, Any]] = []
                    cell_cursor = 0

                    for paragraph in cell.findall("./w:p", DOCX_NS_MAP):
                        paragraph_data = _parse_paragraph_content(paragraph, comments_by_id)
                        paragraph_text = str(paragraph_data.get("text") or "")
                        if not paragraph_text:
                            continue

                        if cell_text_parts:
                            cell_text_parts.append("\n")
                            cell_runs.append({"text": "\n"})
                            cell_cursor += 1

                        cell_text_parts.append(paragraph_text)
                        cell_runs.extend(list(paragraph_data.get("runs") or []))
                        for anchor in list(paragraph_data.get("comment_anchors") or []):
                            mapped = dict(anchor)
                            if isinstance(mapped.get("char_start"), int):
                                mapped["char_start"] = cell_cursor + int(mapped["char_start"])
                            if isinstance(mapped.get("char_end"), int):
                                mapped["char_end"] = cell_cursor + int(mapped["char_end"])
                            cell_anchors.append(mapped)
                        cell_cursor += len(paragraph_text)

                    cell_text = "".join(cell_text_parts).strip()
                    row_cells.append({"text": cell_text, "runs": _merge_adjacent_runs(cell_runs)})
                    row_text_parts.append(cell_text)

                    for anchor in cell_anchors:
                        mapped = dict(anchor)
                        if isinstance(mapped.get("char_start"), int):
                            mapped["char_start"] = row_cursor + int(mapped["char_start"])
                        if isinstance(mapped.get("char_end"), int):
                            mapped["char_end"] = row_cursor + int(mapped["char_end"])
                        row_anchor_items.append(mapped)

                    row_cursor += len(cell_text) + 3  # " | "

                rows_payload.append({"cells": row_cells})
                row_text = " | ".join(row_text_parts)
                row_line_texts.append(row_text)
                for anchor in row_anchor_items:
                    mapped = dict(anchor)
                    if isinstance(mapped.get("char_start"), int):
                        mapped["char_start"] = table_offset + int(mapped["char_start"])
                    if isinstance(mapped.get("char_end"), int):
                        mapped["char_end"] = table_offset + int(mapped["char_end"])
                    table_anchors.append(mapped)
                table_offset += len(row_text) + 1

            table_text = "\n".join(row_line_texts).strip()
            if not table_text and not rows_payload:
                continue

            append_block(
                {
                    "type": "table",
                    "text": table_text,
                    "table": {"rows": rows_payload},
                    "meta": {
                        "row_count": len(rows_payload),
                        "column_count": max((len(row.get("cells") or []) for row in rows_payload), default=0),
                    },
                },
                table_anchors,
            )

    markdown = "\n\n".join(line for line in markdown_lines if line).strip()
    block_text_by_id = {str(block.get("id")): str(block.get("text") or "") for block in blocks}

    docx_comments: List[Dict[str, Any]] = []
    seen_comment_keys = set()
    for anchor in anchor_items:
        comment_id = str(anchor.get("comment_id") or "").strip()
        block_id = str(anchor.get("block_id") or "").strip()
        if not comment_id or not block_id:
            continue
        key = (comment_id, block_id)
        if key in seen_comment_keys:
            continue
        seen_comment_keys.add(key)

        comment_data = comments_by_id.get(comment_id, {})
        char_start = anchor.get("char_start")
        char_end = anchor.get("char_end")
        selected_text = str(anchor.get("selected_text") or "").strip()
        block_text = block_text_by_id.get(block_id, "")
        if (
            not selected_text
            and isinstance(char_start, int)
            and isinstance(char_end, int)
            and 0 <= char_start <= char_end <= len(block_text)
        ):
            selected_text = block_text[char_start:char_end].strip()

        docx_comments.append(
            {
                "comment_id": comment_id,
                "block_id": block_id,
                "char_start": char_start if isinstance(char_start, int) else None,
                "char_end": char_end if isinstance(char_end, int) else None,
                "selected_text": selected_text or None,
                "author": str(comment_data.get("author") or "").strip() or None,
                "initials": str(comment_data.get("initials") or "").strip() or None,
                "date": str(comment_data.get("date") or "").strip() or None,
                "content": str(comment_data.get("content") or "").strip() or "",
            }
        )

    return {
        "normalized_markdown": markdown,
        "blocks_json": blocks,
        "render_json": {
            "format": "rich_doc",
            "block_count": len(blocks),
            "docx_comments": docx_comments,
            "features": {
                "runs": True,
                "tables": True,
                "toc": True,
                "comments": bool(docx_comments),
            },
        },
    }


def _docx_to_payload(raw: bytes) -> Dict[str, Any]:
    parsed = _parse_docx_payload_via_xml(raw)
    if parsed.get("blocks_json") or parsed.get("normalized_markdown"):
        return parsed

    markdown = ""
    blocks: List[Dict[str, Any]] = []
    if DocxDocument:
        try:
            markdown, blocks = _docx_to_markdown_via_python_docx(raw)
        except Exception:
            markdown = ""
            blocks = []

    return {
        "normalized_markdown": markdown,
        "blocks_json": blocks,
        "render_json": {
            "format": "rich_doc",
            "block_count": len(blocks),
            "docx_comments": [],
            "features": {
                "runs": bool(blocks),
                "tables": False,
                "toc": False,
                "comments": False,
            },
        },
    }


def _docx_to_markdown(raw: bytes) -> Tuple[str, List[Dict[str, Any]]]:
    payload = _docx_to_payload(raw)
    return str(payload.get("normalized_markdown") or ""), list(payload.get("blocks_json") or [])


def _docx_to_markdown_via_python_docx(raw: bytes) -> Tuple[str, List[Dict[str, Any]]]:
    doc = DocxDocument(io.BytesIO(raw))  # type: ignore[misc]

    lines: List[str] = []
    blocks: List[Dict[str, Any]] = []
    order = 0
    for paragraph in doc.paragraphs:
        text = (paragraph.text or "").strip()
        if not text:
            continue
        style_name = (getattr(paragraph.style, "name", "") or "").strip()
        style_name_l = style_name.lower()

        block_type = "paragraph"
        meta: Dict[str, Any] = {"style": style_name}
        markdown_line = text

        if style_name_l.startswith("heading"):
            level_match = re.search(r"(\d+)", style_name_l)
            level = int(level_match.group(1)) if level_match else 1
            level = min(max(level, 1), 6)
            markdown_line = f"{'#' * level} {text}"
            block_type = "heading"
            meta["level"] = level
        elif "list bullet" in style_name_l or "list paragraph" in style_name_l:
            markdown_line = f"- {text}"
            block_type = "list_item"
            meta["marker"] = "-"
        elif "list number" in style_name_l:
            markdown_line = f"1. {text}"
            block_type = "list_item"
            meta["marker"] = "1."

        lines.append(markdown_line)
        order += 1
        blocks.append(
            {
                "id": f"blk-{order}",
                "type": block_type,
                "text": text,
                "order": order,
                "meta": meta,
            }
        )

    markdown = "\n\n".join(lines).strip()
    if not blocks:
        return "", []
    return markdown, blocks


def _docx_to_markdown_via_xml(raw: bytes) -> Tuple[str, List[Dict[str, Any]]]:
    try:
        with zipfile.ZipFile(io.BytesIO(raw)) as zf:
            xml_bytes = zf.read("word/document.xml")
    except Exception:
        return "", []

    try:
        root = ET.fromstring(xml_bytes)
    except Exception:
        return "", []

    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    lines: List[str] = []
    blocks: List[Dict[str, Any]] = []
    order = 0

    for paragraph in root.findall(".//w:body/w:p", ns):
        text_parts: List[str] = []
        for node in paragraph.findall(".//w:t", ns):
            if node.text:
                text_parts.append(node.text)
        text = "".join(text_parts).strip()
        if not text:
            continue

        block_type = "paragraph"
        meta: Dict[str, Any] = {}
        markdown_line = text

        p_style = paragraph.find("./w:pPr/w:pStyle", ns)
        style_val = ""
        if p_style is not None:
            style_val = p_style.attrib.get(f"{{{ns['w']}}}val", "").strip()
        style_l = style_val.lower()

        if style_l.startswith("heading"):
            level_match = re.search(r"(\d+)", style_l)
            level = int(level_match.group(1)) if level_match else 1
            level = min(max(level, 1), 6)
            markdown_line = f"{'#' * level} {text}"
            block_type = "heading"
            meta["level"] = level
            meta["style"] = style_val
        else:
            num_pr = paragraph.find("./w:pPr/w:numPr", ns)
            if num_pr is not None:
                level_node = num_pr.find("./w:ilvl", ns)
                level = 0
                if level_node is not None:
                    try:
                        level = int(level_node.attrib.get(f"{{{ns['w']}}}val", "0"))
                    except Exception:
                        level = 0
                indent = "  " * max(level, 0)
                markdown_line = f"{indent}- {text}"
                block_type = "list_item"
                meta["marker"] = "-"
                meta["level"] = level
            elif style_val:
                meta["style"] = style_val

        lines.append(markdown_line)
        order += 1
        blocks.append(
            {
                "id": f"blk-{order}",
                "type": block_type,
                "text": text,
                "order": order,
                "meta": meta,
            }
        )

    markdown = "\n\n".join(lines).strip()
    return markdown, blocks


def _looks_like_binary_docx_dump(text: str) -> bool:
    if not text:
        return False
    if text.startswith("PK") and ("word/" in text or "[Content_Types].xml" in text):
        return True
    control_count = sum(1 for ch in text if ord(ch) < 32 and ch not in "\n\r\t")
    return control_count > max(8, len(text) // 200)


def repair_docx_version_if_needed(db: Session, asset: SddAsset, version: SddAssetVersion) -> bool:
    ext = (version.original_ext or asset.source_ext or "").lower()
    if ext != ".docx":
        return False

    markdown = version.normalized_markdown or ""
    blocks = list(version.blocks_json or []) if isinstance(version.blocks_json, list) else []
    render = version.render_json or {}
    render_format = str(render.get("format") or "").strip().lower()
    has_docx_comment_payload = isinstance(render.get("docx_comments"), list)
    needs_repair = (
        (not blocks)
        or _looks_like_binary_docx_dump(markdown)
        or render_format != "rich_doc"
        or (has_docx_comment_payload is False)
    )

    original_path = (version.original_path or "").strip()
    if not original_path or not os.path.isfile(original_path):
        return False

    try:
        with open(original_path, "rb") as f:
            raw = f.read()
    except Exception:
        return False

    original_docx_valid = _looks_like_docx_bytes(raw)
    needs_repair = needs_repair or (not original_docx_valid)
    if not needs_repair:
        return False

    if not original_docx_valid:
        rebuilt = _build_docx_bytes_from_content(blocks, markdown)
        if rebuilt:
            try:
                with open(original_path, "wb") as f:
                    f.write(rebuilt)
                raw = rebuilt
                original_docx_valid = True
            except Exception:
                pass

    repaired_payload = parse_document_payload(
        version.original_path or asset.source_file_name or "document.docx",
        raw,
    )
    repaired_markdown = str(repaired_payload.get("normalized_markdown") or "")
    repaired_blocks = list(repaired_payload.get("blocks_json") or [])
    repaired_render = dict(repaired_payload.get("render_json") or {})
    if not repaired_markdown and not repaired_blocks:
        return False

    version.normalized_markdown = repaired_markdown
    version.blocks_json = repaired_blocks
    version.render_json = repaired_render or {
        "format": "rich_doc",
        "block_count": len(repaired_blocks),
        "docx_comments": [],
    }

    if asset.active_version_id == version.id:
        asset.content_text = repaired_markdown
        asset.content_json = {
            "active_version_no": version.version_no,
            "block_count": len(repaired_blocks),
        }

    db.flush()
    return True


def parse_document_payload(file_name: str, raw: bytes) -> Dict[str, Any]:
    ext, mime = _guess_ext_and_mime(file_name)
    render_json: Dict[str, Any]
    if ext in {".md", ".markdown"}:
        markdown = _decode_text_bytes(raw)
        blocks = _markdown_to_blocks(markdown)
        render_json = {
            "format": "markdown",
            "block_count": len(blocks),
        }
    elif ext == ".txt":
        markdown = _plain_text_to_markdown(_decode_text_bytes(raw))
        blocks = _markdown_to_blocks(markdown)
        render_json = {
            "format": "markdown",
            "block_count": len(blocks),
        }
    elif ext == ".docx":
        docx_payload = _docx_to_payload(raw)
        markdown = str(docx_payload.get("normalized_markdown") or "")
        blocks = list(docx_payload.get("blocks_json") or [])
        render_json = dict(docx_payload.get("render_json") or {})
    elif ext == ".pdf":
        markdown = ""
        blocks = []
        render_json = {
            "format": "markdown",
            "block_count": 0,
        }
    else:
        markdown = _plain_text_to_markdown(_decode_text_bytes(raw))
        blocks = _markdown_to_blocks(markdown)
        render_json = {
            "format": "markdown",
            "block_count": len(blocks),
        }

    return {
        "source_file_name": file_name,
        "source_ext": ext,
        "source_mime": mime,
        "normalized_markdown": markdown,
        "blocks_json": blocks,
        "render_json": render_json,
    }


def can_inline_review(ext: Optional[str]) -> bool:
    return (ext or "").lower() in SUPPORTED_INLINE_REVIEW_EXTENSIONS


def get_spec_asset_by_task(db: Session, task_id: str) -> Optional[SddAsset]:
    return (
        db.query(SddAsset)
        .filter(
            SddAsset.task_id == task_id,
            SddAsset.asset_type == AssetType.SPEC,
        )
        .order_by(SddAsset.created_at.asc())
        .first()
    )


def get_asset_by_id(db: Session, workspace_id: str, asset_id: str) -> Optional[SddAsset]:
    return (
        db.query(SddAsset)
        .filter(SddAsset.id == asset_id, SddAsset.workspace_id == workspace_id)
        .first()
    )


def list_asset_versions(db: Session, asset_id: str) -> List[SddAssetVersion]:
    return (
        db.query(SddAssetVersion)
        .filter(SddAssetVersion.asset_id == asset_id)
        .order_by(SddAssetVersion.version_no.desc(), SddAssetVersion.created_at.desc())
        .all()
    )


def get_asset_version(db: Session, asset_id: str, version_id: str) -> Optional[SddAssetVersion]:
    return (
        db.query(SddAssetVersion)
        .filter(
            SddAssetVersion.asset_id == asset_id,
            SddAssetVersion.id == version_id,
        )
        .first()
    )


def _next_version_no(db: Session, asset_id: str) -> int:
    max_no = (
        db.query(func.max(SddAssetVersion.version_no))
        .filter(SddAssetVersion.asset_id == asset_id)
        .scalar()
    )
    return int(max_no or 0) + 1


def _version_dir(task: SddTask, asset: SddAsset, version_no: int) -> str:
    path = os.path.join(
        _task_assets_root(task),
        asset.id,
        f"v{version_no}",
    )
    os.makedirs(path, exist_ok=True)
    return path


def _write_original_file(
    task: SddTask,
    asset: SddAsset,
    version_no: int,
    file_name: str,
    file_content: bytes,
) -> str:
    storage_dir = _version_dir(task, asset, version_no)
    safe_name = _normalize_filename(file_name)
    target_path = os.path.abspath(os.path.join(storage_dir, safe_name))
    with open(target_path, "wb") as f:
        f.write(file_content)
    return target_path


def _looks_like_docx_bytes(raw: Optional[bytes]) -> bool:
    if not raw or len(raw) < 4:
        return False
    if not raw.startswith(b"PK"):
        return False
    try:
        with zipfile.ZipFile(io.BytesIO(raw)) as zf:
            names = set(zf.namelist())
    except Exception:
        return False
    if "[Content_Types].xml" not in names:
        return False
    return any(name.startswith("word/") for name in names)


def _extract_block_text_for_export(block: Dict[str, Any]) -> str:
    text = str(block.get("text") or "").strip()
    if text:
        return text

    runs = block.get("runs")
    if isinstance(runs, list):
        merged = "".join(str(item.get("text") or "") for item in runs if isinstance(item, dict)).strip()
        if merged:
            return merged

    cells = block.get("cells")
    if isinstance(cells, list):
        row_chunks: List[str] = []
        for row in cells:
            if not isinstance(row, list):
                continue
            cell_texts: List[str] = []
            for cell in row:
                if isinstance(cell, dict):
                    value = str(cell.get("text") or "").strip()
                    if value:
                        cell_texts.append(value)
            if cell_texts:
                row_chunks.append(" | ".join(cell_texts))
        if row_chunks:
            return "\n".join(row_chunks)

    return ""


def _normalize_docx_export_paragraphs(
    blocks: List[Dict[str, Any]],
    markdown: str,
) -> List[Dict[str, Any]]:
    source_blocks: List[Dict[str, Any]] = [item for item in (blocks or []) if isinstance(item, dict)]
    if not source_blocks and str(markdown or "").strip():
        source_blocks = _markdown_to_blocks(str(markdown or ""))

    paragraphs: List[Dict[str, Any]] = []
    for block in source_blocks:
        text = _extract_block_text_for_export(block)
        if not text:
            continue
        paragraphs.append(
            {
                "type": str(block.get("type") or "paragraph").strip().lower(),
                "text": text,
                "meta": block.get("meta") if isinstance(block.get("meta"), dict) else {},
            }
        )

    if paragraphs:
        return paragraphs

    fallback = str(markdown or "").strip()
    if not fallback:
        return [{"type": "paragraph", "text": "", "meta": {}}]
    return [{"type": "paragraph", "text": fallback, "meta": {}}]


def _format_docx_export_line(paragraph: Dict[str, Any]) -> str:
    text = str(paragraph.get("text") or "")
    p_type = str(paragraph.get("type") or "paragraph").strip().lower()
    meta = paragraph.get("meta") if isinstance(paragraph.get("meta"), dict) else {}
    if p_type == "list_item":
        marker = str(meta.get("marker") or "-").strip() or "-"
        return f"{marker} {text}".strip()
    return text


def _build_minimal_docx_bytes(paragraphs: List[Dict[str, Any]]) -> bytes:
    content_types_xml = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
  <Default Extension="xml" ContentType="application/xml"/>
  <Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
</Types>"""
    rels_xml = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>
</Relationships>"""

    body_parts: List[str] = []
    for paragraph in paragraphs:
        text = _format_docx_export_line(paragraph).replace("\r\n", "\n").replace("\r", "\n")
        lines = text.split("\n") if text else [""]
        for line in lines:
            if line:
                body_parts.append(
                    f'<w:p><w:r><w:t xml:space="preserve">{html.escape(line)}</w:t></w:r></w:p>'
                )
            else:
                body_parts.append("<w:p/>")
    if not body_parts:
        body_parts.append("<w:p/>")

    document_xml = (
        """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>"""
        """<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">"""
        f"""<w:body>{''.join(body_parts)}<w:sectPr/></w:body>"""
        """</w:document>"""
    )

    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, mode="w", compression=zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("[Content_Types].xml", content_types_xml)
        zf.writestr("_rels/.rels", rels_xml)
        zf.writestr("word/document.xml", document_xml)
    return buffer.getvalue()


def _build_docx_bytes_from_content(
    blocks: List[Dict[str, Any]],
    markdown: str,
) -> Optional[bytes]:
    paragraphs = _normalize_docx_export_paragraphs(blocks, markdown)

    if DocxDocument:
        try:
            doc = DocxDocument()
            if len(doc.paragraphs) == 1 and not str(doc.paragraphs[0].text or "").strip():
                element = doc.paragraphs[0]._element
                element.getparent().remove(element)

            for paragraph in paragraphs:
                text = str(paragraph.get("text") or "")
                p_type = str(paragraph.get("type") or "paragraph").strip().lower()
                meta = paragraph.get("meta") if isinstance(paragraph.get("meta"), dict) else {}
                if p_type == "heading":
                    level = _safe_int(meta.get("level"), default=1)
                    level = min(max(level, 1), 6)
                    doc.add_heading(text, level=level)
                    continue
                if p_type == "list_item":
                    marker = str(meta.get("marker") or "-").strip()
                    style = "List Number" if re.match(r"^\d+\.$", marker) else "List Bullet"
                    para = doc.add_paragraph(style=style)
                    para.add_run(text)
                    continue
                doc.add_paragraph(text)

            if not doc.paragraphs:
                doc.add_paragraph("")
            output = io.BytesIO()
            doc.save(output)
            return output.getvalue()
        except Exception:
            pass

    try:
        return _build_minimal_docx_bytes(paragraphs)
    except Exception:
        return None


def _upsert_spec_asset(
    db: Session,
    task: SddTask,
    creator_id: str,
    file_name: str,
    payload: Dict[str, Any],
) -> SddAsset:
    asset = get_spec_asset_by_task(db, task.id)
    if asset:
        asset.name = file_name
        asset.source_file_name = payload.get("source_file_name")
        asset.source_ext = payload.get("source_ext")
        asset.source_mime = payload.get("source_mime")
        return asset

    asset = SddAsset(
        task_id=task.id,
        workspace_id=task.workspace_id,
        creator_id=creator_id,
        asset_type=AssetType.SPEC,
        name=file_name,
        source_file_name=payload.get("source_file_name"),
        source_ext=payload.get("source_ext"),
        source_mime=payload.get("source_mime"),
    )
    db.add(asset)
    db.flush()
    return asset


def create_asset_version_from_upload(
    db: Session,
    task: SddTask,
    *,
    creator_id: str,
    file_name: str,
    file_content: bytes,
    change_note: Optional[str] = None,
) -> Tuple[SddAsset, SddAssetVersion]:
    payload = parse_document_payload(file_name, file_content)
    asset = _upsert_spec_asset(db, task, creator_id, file_name, payload)
    base_version_id = asset.active_version_id
    version_no = _next_version_no(db, asset.id)

    original_path = _write_original_file(task, asset, version_no, file_name, file_content)
    version = SddAssetVersion(
        asset_id=asset.id,
        version_no=version_no,
        base_version_id=base_version_id,
        original_path=original_path,
        original_ext=payload.get("source_ext"),
        original_mime=payload.get("source_mime"),
        normalized_markdown=payload.get("normalized_markdown"),
        blocks_json=payload.get("blocks_json"),
        render_json=payload.get("render_json"),
        change_note=change_note,
        created_by=creator_id,
    )
    db.add(version)
    db.flush()

    asset.active_version_id = version.id
    asset.content_text = payload.get("normalized_markdown")
    asset.content_json = {
        "active_version_no": version_no,
        "block_count": len(payload.get("blocks_json") or []),
    }
    asset.source_file_name = payload.get("source_file_name")
    asset.source_ext = payload.get("source_ext")
    asset.source_mime = payload.get("source_mime")

    return asset, version


def get_diagnosis_doc_asset_by_task_and_name(
    db: Session, task_id: str, file_name: str
) -> Optional[SddAsset]:
    return (
        db.query(SddAsset)
        .filter(
            SddAsset.task_id == task_id,
            SddAsset.asset_type == AssetType.DIAGNOSIS_DOC,
            SddAsset.name == file_name,
        )
        .order_by(SddAsset.created_at.asc())
        .first()
    )


def create_diagnosis_doc_asset_version(
    db: Session,
    task: SddTask,
    *,
    creator_id: str,
    file_name: str,
    file_content: bytes,
    change_note: Optional[str] = None,
) -> Tuple[SddAsset, SddAssetVersion, str]:
    """问题定位任务：上传需求/日志等辅助文档。

    - 按「任务 + 文件名」复用同一 DIAGNOSIS_DOC 资产（重复上传生成新版本）；
    - 原始文件同时写入任务 CLI 工作区 `.sdd/diagnosis/` 目录，AI 会话可直接读取；
    - 返回 (asset, version, cli_path)。
    """
    safe_name = _normalize_filename(file_name)
    payload = parse_document_payload(safe_name, file_content)
    asset = get_diagnosis_doc_asset_by_task_and_name(db, task.id, safe_name)
    if asset is None:
        asset = SddAsset(
            task_id=task.id,
            workspace_id=task.workspace_id,
            creator_id=creator_id,
            asset_type=AssetType.DIAGNOSIS_DOC,
            name=safe_name,
            source_file_name=payload.get("source_file_name"),
            source_ext=payload.get("source_ext"),
            source_mime=payload.get("source_mime"),
        )
        db.add(asset)
        db.flush()

    base_version_id = asset.active_version_id
    version_no = _next_version_no(db, asset.id)
    original_path = _write_original_file(task, asset, version_no, safe_name, file_content)
    version = SddAssetVersion(
        asset_id=asset.id,
        version_no=version_no,
        base_version_id=base_version_id,
        original_path=original_path,
        original_ext=payload.get("source_ext"),
        original_mime=payload.get("source_mime"),
        normalized_markdown=payload.get("normalized_markdown"),
        blocks_json=payload.get("blocks_json"),
        render_json=payload.get("render_json"),
        change_note=change_note,
        created_by=creator_id,
    )
    db.add(version)
    db.flush()

    asset.active_version_id = version.id
    asset.content_text = payload.get("normalized_markdown")
    asset.content_json = {
        "active_version_no": version_no,
        "block_count": len(payload.get("blocks_json") or []),
    }
    asset.source_file_name = payload.get("source_file_name")
    asset.source_ext = payload.get("source_ext")
    asset.source_mime = payload.get("source_mime")
    asset.name = safe_name

    # CLI 工作区副本：.sdd/diagnosis/<file>（AI 会话可直接读取辅助文档）
    base_dir = str(getattr(task, "project_path", "") or "").strip() or os.getcwd()
    cli_dir = os.path.abspath(os.path.join(base_dir, ".sdd", "diagnosis"))
    os.makedirs(cli_dir, exist_ok=True)
    cli_path = os.path.join(cli_dir, safe_name)
    with open(cli_path, "wb") as f:
        f.write(file_content)

    db.flush()
    return asset, version, cli_path


def create_asset_version_from_normalized_content(
    db: Session,
    asset: SddAsset,
    *,
    creator_id: str,
    normalized_markdown: str,
    blocks_json: Optional[List[Dict[str, Any]]] = None,
    change_note: Optional[str] = None,
    base_version_id: Optional[str] = None,
    output_ext: Optional[str] = None,
    output_mime: Optional[str] = None,
    output_file_bytes: Optional[bytes] = None,
    output_file_name: Optional[str] = None,
) -> SddAssetVersion:
    task = _task_for_asset(db, asset)
    version_no = _next_version_no(db, asset.id)
    ext = (output_ext or asset.source_ext or ".md").lower()
    mime = output_mime or asset.source_mime or "text/markdown"
    effective_markdown = normalized_markdown
    blocks = blocks_json if blocks_json is not None else _markdown_to_blocks(effective_markdown)
    render_json: Dict[str, Any] = {"format": "markdown", "block_count": len(blocks)}
    filename = output_file_name or (asset.source_file_name or f"spec-v{version_no}{ext}")
    file_name = _normalize_filename(filename, fallback_ext=ext or ".md")
    file_bytes = output_file_bytes
    if file_bytes is None:
        file_bytes = effective_markdown.encode("utf-8")

    if ext == ".docx" and not _looks_like_docx_bytes(file_bytes):
        rebuilt_docx = _build_docx_bytes_from_content(blocks, effective_markdown)
        if rebuilt_docx:
            file_bytes = rebuilt_docx
        mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    if ext == ".docx" and file_bytes:
        reparsed = parse_document_payload(file_name, file_bytes)
        reparsed_markdown = str(reparsed.get("normalized_markdown") or "")
        reparsed_blocks = list(reparsed.get("blocks_json") or [])
        reparsed_render = dict(reparsed.get("render_json") or {})
        if reparsed_blocks or reparsed_markdown:
            if reparsed_markdown:
                effective_markdown = reparsed_markdown
            if reparsed_blocks:
                blocks = reparsed_blocks
            render_json = reparsed_render or {"format": "rich_doc", "block_count": len(blocks), "docx_comments": []}

    original_path = _write_original_file(task, asset, version_no, file_name, file_bytes)

    version = SddAssetVersion(
        asset_id=asset.id,
        version_no=version_no,
        base_version_id=base_version_id or asset.active_version_id,
        original_path=original_path,
        original_ext=ext,
        original_mime=mime,
        normalized_markdown=effective_markdown,
        blocks_json=blocks,
        render_json=render_json,
        change_note=change_note,
        created_by=creator_id,
    )
    db.add(version)
    db.flush()

    asset.active_version_id = version.id
    asset.content_text = effective_markdown
    asset.content_json = {
        "active_version_no": version_no,
        "block_count": len(blocks),
    }
    asset.source_file_name = file_name
    asset.source_ext = ext
    asset.source_mime = mime
    return version


def create_task_asset_version_from_bytes(
    db: Session,
    task: SddTask,
    *,
    creator_id: str,
    asset_type: AssetType,
    asset_name: str,
    file_name: str,
    file_content: bytes,
    content_text: Optional[str] = None,
    content_json: Optional[Dict[str, Any]] = None,
    change_note: Optional[str] = None,
    source_ext: Optional[str] = None,
    source_mime: Optional[str] = None,
) -> Tuple[SddAsset, SddAssetVersion]:
    ext, guessed_mime = _guess_ext_and_mime(file_name)
    effective_ext = source_ext if source_ext is not None else ext
    effective_mime = source_mime or guessed_mime
    asset = SddAsset(
        task_id=task.id,
        workspace_id=task.workspace_id,
        creator_id=creator_id,
        asset_type=asset_type,
        name=asset_name,
        content_text=content_text,
        content_json=content_json or {},
        source_file_name=file_name,
        source_ext=effective_ext,
        source_mime=effective_mime,
    )
    db.add(asset)
    db.flush()

    version_no = _next_version_no(db, asset.id)
    original_path = _write_original_file(task, asset, version_no, file_name, file_content)
    version = SddAssetVersion(
        asset_id=asset.id,
        version_no=version_no,
        original_path=original_path,
        original_ext=effective_ext,
        original_mime=effective_mime,
        normalized_markdown=content_text,
        blocks_json=[],
        render_json={
            "format": "artifact",
            "source_file_name": file_name,
            "size_bytes": len(file_content or b""),
        },
        change_note=change_note,
        created_by=creator_id,
    )
    db.add(version)
    db.flush()

    asset.active_version_id = version.id
    return asset, version


def ensure_spec_asset_backfilled(db: Session, task: SddTask) -> Optional[SddAsset]:
    """
    For legacy tasks, lazily backfill SPEC asset/version from task.spec_doc_path.
    """
    existing = get_spec_asset_by_task(db, task.id)
    if existing:
        return existing

    spec_path = (task.spec_doc_path or "").strip()
    if not spec_path:
        return None
    abs_spec_path = os.path.abspath(spec_path)
    if not os.path.exists(abs_spec_path) or not os.path.isfile(abs_spec_path):
        return None

    try:
        with open(abs_spec_path, "rb") as f:
            raw = f.read()
    except Exception:
        return None

    file_name = os.path.basename(abs_spec_path)
    asset, _version = create_asset_version_from_upload(
        db,
        task,
        creator_id=task.creator_id,
        file_name=file_name,
        file_content=raw,
        change_note="Initial version (backfilled from task.spec_doc_path)",
    )
    return asset


def ensure_asset_has_version(db: Session, asset: SddAsset) -> Optional[SddAssetVersion]:
    if asset.active_version_id:
        active = get_asset_version(db, asset.id, asset.active_version_id)
        if active:
            return active

    versions = list_asset_versions(db, asset.id)
    if versions:
        newest = versions[0]
        asset.active_version_id = newest.id
        db.flush()
        return newest

    text = (asset.content_text or "").strip()
    if not text:
        return None

    source_name = asset.source_file_name or "legacy-spec.md"
    version = create_asset_version_from_normalized_content(
        db,
        asset,
        creator_id=asset.creator_id,
        normalized_markdown=text,
        change_note="Initial version (backfilled from legacy asset content)",
        output_ext=asset.source_ext or ".md",
        output_mime=asset.source_mime or "text/markdown",
        output_file_name=source_name,
    )
    return version
